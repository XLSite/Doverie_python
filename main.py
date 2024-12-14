import os
import docx
tree = os.walk('docs')
def getdocx(document):
    table = document.tables[0]
    Project_cell = table.rows[2].cells[2]
    paragraph = Project_cell.paragraphs[0]
    Project = paragraph.text
    return Project

for i in tree:
    getdocx()



if __name__ == "__main__":
    paragraphs = []
    f = os.walk(tree).next()
    for filename in f:
        file_name = os.path.join(tree, filename)
        document = docx.Document(file_name)
        para = getdocx(document)
        paragraphs.append(para)

    print(paragraphs)

